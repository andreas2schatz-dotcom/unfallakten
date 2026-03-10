"""
word_generator.py – Forderungsschreiben aus Vorlage + Datenbank befüllen

Ablauf:
  1. Vorlage laden (Forderungsschreiben_Vorlage.docx)
  2. Alle {{PLATZHALTER}} durch echte Werte ersetzen
  3. Datei unter Aktenzeichen speichern und Pfad zurückgeben
"""

import os
import re
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn


# ──────────────────────────────────────────────────────────────
#  KONFIGURATION
# ──────────────────────────────────────────────────────────────
VORLAGE_PFAD   = Path(__file__).parent / "templates" / "Forderungsschreiben_Vorlage.docx"
AUSGABE_ORDNER = Path(__file__).parent / "ausgabe"
AUSGABE_ORDNER.mkdir(exist_ok=True)

# Kanzlei-Stammdaten (alternativ aus .env laden)
KANZLEI = {
    "ANWALT_NAME":    os.environ.get("KANZLEI_NAME",    "Rechtsanwalt Max Mustermann"),
    "ANWALT_STRASSE": os.environ.get("KANZLEI_STRASSE", "Musterstraße 1"),
    "ANWALT_PLZ_ORT": os.environ.get("KANZLEI_PLZ_ORT", "12345 Musterstadt"),
    "ANWALT_TEL":     os.environ.get("KANZLEI_TEL",     "0123 456789"),
    "ANWALT_EMAIL":   os.environ.get("KANZLEI_EMAIL",   "kanzlei@muster.de"),
    "IBAN":           os.environ.get("KANZLEI_IBAN",    "DE12 3456 7890 1234 5678 90"),
    "BIC":            os.environ.get("KANZLEI_BIC",     "MUSTDE12XXX"),
}


# ──────────────────────────────────────────────────────────────
#  HILFSFUNKTIONEN
# ──────────────────────────────────────────────────────────────
def _eur(wert: Optional[float]) -> str:
    """Formatiert einen Float als deutschen EUR-Betrag, z.B. 6200.5 → '6.200,50'"""
    if wert is None or wert == 0:
        return "0,00"
    return f"{wert:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _datum(d: Optional[date]) -> str:
    if d is None:
        return ""
    if isinstance(d, str):
        try:
            d = datetime.strptime(d, "%Y-%m-%d").date()
        except ValueError:
            return d
    return d.strftime("%d.%m.%Y")


def _replace_in_paragraph(paragraph, platzhalter: dict):
    """Ersetzt Platzhalter in einem Absatz, auch wenn der Text auf mehrere Runs verteilt ist."""
    # Gesamttext des Absatzes zusammenbauen
    full_text = "".join(run.text for run in paragraph.runs)

    # Prüfen ob überhaupt ein Platzhalter enthalten ist
    if "{{" not in full_text:
        return

    # Alle Ersetzungen durchführen
    new_text = full_text
    for key, value in platzhalter.items():
        new_text = new_text.replace(f"{{{{{key}}}}}", str(value) if value is not None else "")

    # Wenn sich etwas geändert hat: ersten Run setzen, Rest leeren
    if new_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


def _replace_in_table(table, platzhalter: dict):
    """Alle Zellen einer Tabelle durchsuchen."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, platzhalter)


def _replace_in_header_footer(section, platzhalter: dict):
    """Kopf- und Fußzeilen ersetzen."""
    for hf in [section.header, section.footer]:
        if hf:
            for paragraph in hf.paragraphs:
                _replace_in_paragraph(paragraph, platzhalter)
            for table in hf.tables:
                _replace_in_table(table, platzhalter)


# ──────────────────────────────────────────────────────────────
#  HAUPTFUNKTION
# ──────────────────────────────────────────────────────────────
def generiere_forderungsschreiben(
    unfall_data: dict,
    ausgabe_dateiname: Optional[str] = None
) -> str:
    """
    Erzeugt ein Forderungsschreiben als .docx-Datei.

    Args:
        unfall_data: Dictionary mit allen Falldaten (aus der Datenbank zusammengestellt).
                     Struktur: siehe _baue_platzhalter() unten.
        ausgabe_dateiname: Optionaler Dateiname, sonst automatisch aus AZ generiert.

    Returns:
        Absoluter Pfad zur erzeugten .docx-Datei.

    Raises:
        FileNotFoundError: Wenn die Vorlage nicht gefunden wird.
    """
    if not VORLAGE_PFAD.exists():
        raise FileNotFoundError(
            f"Vorlage nicht gefunden: {VORLAGE_PFAD}\n"
            "Bitte Forderungsschreiben_Vorlage.docx in den Ordner 'templates/' legen."
        )

    # Platzhalter-Dictionary aufbauen
    platzhalter = _baue_platzhalter(unfall_data)

    # Vorlage laden
    doc = Document(VORLAGE_PFAD)

    # ── Alle Textbereiche ersetzen ──────────────────────────
    # 1. Normale Absätze
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, platzhalter)

    # 2. Tabellen
    for table in doc.tables:
        _replace_in_table(table, platzhalter)

    # 3. Kopf- und Fußzeilen
    for section in doc.sections:
        _replace_in_header_footer(section, platzhalter)

    # ── Datei speichern ─────────────────────────────────────
    az_clean = re.sub(r"[^\w\-]", "_", unfall_data.get("aktenzeichen", "unbekannt"))
    dateiname = ausgabe_dateiname or f"Forderungsschreiben_{az_clean}_{datetime.now():%Y%m%d}.docx"
    ausgabe_pfad = AUSGABE_ORDNER / dateiname

    doc.save(ausgabe_pfad)
    print(f"✅  Forderungsschreiben gespeichert: {ausgabe_pfad}")

    return str(ausgabe_pfad)


# ──────────────────────────────────────────────────────────────
#  PLATZHALTER AUFBAUEN
# ──────────────────────────────────────────────────────────────
def _baue_platzhalter(d: dict) -> dict:
    """
    Baut das vollständige Platzhalter-Dict aus den Falldaten auf.
    Berechnet Gesamtforderung automatisch.

    Erwartete Schlüssel in d (alle optional, werden mit "" befüllt):
      aktenzeichen, unfalldatum, unfallort,
      mandant_name, gegner_name,
      kennzeichen_mandant, kennzeichen_gegner,
      fahrzeug_marke,
      versicherung_name, versicherung_strasse, versicherung_plz_ort,
      schadennummer,
      betrag_reparatur, betrag_wertminderung, betrag_mietwagen,
      betrag_schmerzensgeld, betrag_gutachter, betrag_auslagen,
      betrag_anwaltskosten,
      zahlungsfrist  (date oder str "DD.MM.YYYY")
    """
    # Einzelbeträge
    reparatur      = float(d.get("betrag_reparatur",      0) or 0)
    wertminderung  = float(d.get("betrag_wertminderung",  0) or 0)
    mietwagen      = float(d.get("betrag_mietwagen",      0) or 0)
    schmerzensgeld = float(d.get("betrag_schmerzensgeld", 0) or 0)
    gutachter      = float(d.get("betrag_gutachter",      0) or 0)
    auslagen       = float(d.get("betrag_auslagen",       0) or 0)
    anwaltskosten  = float(d.get("betrag_anwaltskosten",  0) or 0)
    gesamt         = reparatur + wertminderung + mietwagen + schmerzensgeld + gutachter + auslagen + anwaltskosten

    # Zahlungsfrist
    frist = d.get("zahlungsfrist")
    if isinstance(frist, date):
        frist_str = _datum(frist)
    elif isinstance(frist, str) and len(frist) == 10 and "-" in frist:
        frist_str = _datum(frist)
    else:
        frist_str = str(frist) if frist else ""

    return {
        # Kanzlei
        **KANZLEI,

        # Metadaten
        "DATUM":            _datum(date.today()),
        "AKTENZEICHEN":     d.get("aktenzeichen", ""),
        "SCHADENNUMMER":    d.get("schadennummer", ""),

        # Versicherung
        "VERSICHERUNG_NAME":    d.get("versicherung_name", ""),
        "VERSICHERUNG_STRASSE": d.get("versicherung_strasse", ""),
        "VERSICHERUNG_PLZ_ORT": d.get("versicherung_plz_ort", ""),

        # Unfall
        "UNFALLDATUM":      _datum(d.get("unfalldatum")),
        "UNFALLORT":        d.get("unfallort", ""),

        # Beteiligte
        "MANDANT_NAME":         d.get("mandant_name", ""),
        "GEGNER_NAME":          d.get("gegner_name", ""),
        "KENNZEICHEN_MANDANT":  d.get("kennzeichen_mandant", ""),
        "KENNZEICHEN_GEGNER":   d.get("kennzeichen_gegner", ""),

        # Fahrzeug
        "FAHRZEUG_MARKE":   d.get("fahrzeug_marke", ""),

        # Beträge
        "REPARATURKOSTEN":  _eur(reparatur),
        "WERTMINDERUNG":    _eur(wertminderung),
        "MIETWAGENKOSTEN":  _eur(mietwagen),
        "SCHMERZENSGELD":   _eur(schmerzensgeld),
        "GUTACHTERKOSTEN":  _eur(gutachter),
        "AUSLAGEN":         _eur(auslagen),
        "ANWALTSKOSTEN":    _eur(anwaltskosten),
        "SUMME_GESAMT":     _eur(gesamt),

        # Frist
        "ZAHLUNGSFRIST":    frist_str,

        # Sonstige
        "SEITE":            "1",
    }


# ──────────────────────────────────────────────────────────────
#  HILFSFUNKTION FÜR FASTAPI-ROUTER
# ──────────────────────────────────────────────────────────────
def lade_falldaten_fuer_schreiben(db_session, unfall_id: int) -> dict:
    """
    Liest alle nötigen Daten für ein Forderungsschreiben direkt aus der DB.
    Gibt ein dict zurück, das direkt an generiere_forderungsschreiben() übergeben wird.
    """
    from database import Unfall, Beteiligter, Fahrzeug, Versicherung, Forderung

    unfall = db_session.get(Unfall, unfall_id)
    if not unfall:
        raise ValueError(f"Unfall ID {unfall_id} nicht gefunden.")

    # Mandant
    mandant = db_session.query(Beteiligter).filter_by(
        unfall_id=unfall_id, rolle="mandant"
    ).first()

    # Gegner
    gegner = db_session.query(Beteiligter).filter_by(
        unfall_id=unfall_id, rolle="gegner"
    ).first()

    # Fahrzeug des Mandanten
    fzg = db_session.query(Fahrzeug).filter_by(
        unfall_id=unfall_id, rolle="mandant"
    ).first()

    # Haftpflichtversicherung
    vers = db_session.query(Versicherung).filter_by(
        unfall_id=unfall_id, versicherung_typ="haftpflicht"
    ).first()

    # Letzte Forderung
    forderung = db_session.query(Forderung).filter_by(
        unfall_id=unfall_id
    ).order_by(Forderung.forderungsdatum.desc()).first()

    def _name(b) -> str:
        if not b:
            return ""
        return f"{b.vorname or ''} {b.nachname}".strip()

    return {
        "aktenzeichen":         unfall.aktenzeichen,
        "unfalldatum":          unfall.unfalldatum,
        "unfallort":            unfall.unfallort,

        "mandant_name":         _name(mandant),
        "gegner_name":          _name(gegner),
        "kennzeichen_mandant":  fzg.kennzeichen if fzg else "",
        "kennzeichen_gegner":   "",   # ggf. Gegner-Fahrzeug abfragen

        "fahrzeug_marke":       f"{fzg.marke or ''} {fzg.modell or ''}".strip() if fzg else "",

        "versicherung_name":    vers.versicherung_name if vers else "",
        "versicherung_strasse": vers.strasse if vers else "",
        "versicherung_plz_ort": f"{vers.plz or ''} {vers.ort or ''}".strip() if vers else "",
        "schadennummer":        vers.schadennummer if vers else "",

        "betrag_reparatur":     forderung.betrag_reparatur      if forderung else 0,
        "betrag_wertminderung": forderung.betrag_wertminderung  if forderung else 0,
        "betrag_mietwagen":     forderung.betrag_mietwagen      if forderung else 0,
        "betrag_schmerzensgeld":forderung.betrag_schmerzensgeld if forderung else 0,
        "betrag_gutachter":     forderung.betrag_gutachter      if forderung else 0,
        "betrag_auslagen":      forderung.betrag_auslagen       if forderung else 0,
        "betrag_anwaltskosten": forderung.betrag_anwaltskosten  if forderung else 0,
        "zahlungsfrist":        forderung.frist_datum           if forderung else None,
    }
