"""
Modul 5 – Generator: Forderungsschreiben
==========================================
Erzeugt ein professionelles Forderungsschreiben an den Versicherer
aus den Aktendaten.

Struktur:
  1. Briefkopf (Kanzlei)
  2. Adressblock (Versicherer)
  3. Betreff mit Aktenzeichen + Unfalldatum
  4. Anspruchsgrundlage (Fließtext)
  5. Schadenstabelle mit allen Positionen
  6. Gesamtforderung (fett/Navy)
  7. Zahlungsaufforderung mit Frist
  8. Unterschriftsblock
  9. Fußzeile mit Aktenzeichen + Seitenzahl
"""

import io
import logging
from datetime import date, timedelta
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styling import (
    erstelle_dokument, fuege_briefkopf_ein, fuege_adressblock_ein,
    fuege_abschnittstitel_ein, erstelle_positions_tabelle,
    fuege_summenzeile_ein, fuege_fusszeile_ein,
    fmt_euro, fmt_datum, NAVY, GRAU, SCHWARZ, SCHRIFT_TEXT
)

logger = logging.getLogger(__name__)


def generiere_forderungsschreiben(akte_daten: dict) -> bytes:
    """
    Erstellt ein Forderungsschreiben als Word-Dokument.

    Args:
        akte_daten: Dict mit:
            - akte:          Unfallakte (id, aktenzeichen, unfalldatum, unfallort, haftungsquote)
            - mandant:       Beteiligter mit rolle='mandant' (optional)
            - gegner:        Beteiligter mit rolle='gegner' (optional)
            - schaden:       Schadenposition (optional)
            - regulierungen: Liste von Regulierungen (optional)

    Returns:
        Word-Dokument als Bytes
    """
    doc = erstelle_dokument()
    akte      = akte_daten.get("akte", {})
    mandant   = akte_daten.get("mandant")
    gegner    = akte_daten.get("gegner")
    schaden   = akte_daten.get("schaden")
    kanzlei   = akte_daten.get("kanzlei")

    az        = akte.get("aktenzeichen", "–")
    unfalldatum = fmt_datum(akte.get("unfalldatum", ""))
    unfallort = akte.get("unfallort", "–")
    haftung   = float(akte.get("haftungsquote") or 100)

    # Datum heute + 14-Tage-Frist
    heute = date.today().strftime("%d.%m.%Y")
    frist = (date.today() + timedelta(days=14)).strftime("%d.%m.%Y")

    # ── 1. Briefkopf ─────────────────────────────────────────────────────────
    fuege_briefkopf_ein(doc, kanzlei)

    # ── 2. Empfänger (Versicherung des Gegners) ───────────────────────────────
    empfaenger = _erstelle_empfaenger(gegner)
    betreff = (
        f"Verkehrsunfall vom {unfalldatum}"
        + (f" in {unfallort}" if unfallort and unfallort != "–" else "")
        + f"\nIhr Zeichen: {gegner.get('schaden_nr', '–') if gegner else '–'}"
        + f" | Unser Zeichen: {az}"
    )

    fuege_adressblock_ein(
        doc,
        empfaenger=empfaenger,
        betreff=betreff,
        aktenzeichen=az,
        datum=heute,
    )

    # ── 3. Anspruchsgrundlage ─────────────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Schadensersatzanspruch")

    mandant_name = _vollname(mandant) if mandant else "unser Mandant"
    gegner_name  = _vollname(gegner)  if gegner  else "Ihr Versicherungsnehmer"

    fliesstext = (
        f"In vorbezeichneter Angelegenheit zeigen wir an, dass wir {mandant_name} "
        f"rechtsanwaltlich vertreten.\n\n"
        f"Am {unfalldatum} ereignete sich"
        + (f" in {unfallort}" if unfallort and unfallort != "–" else "")
        + f" ein Verkehrsunfall, an dem {mandant_name} und "
        f"{gegner_name} beteiligt waren. Die alleinige Haftung "
        + ("für den entstandenen Schaden liegt bei Ihrem Versicherungsnehmer."
           if haftung >= 100 else
           f"ist mit {haftung:.0f}% quotiert.")
        + "\n\n"
        f"Wir machen hiermit nachfolgende Schadensersatzansprüche "
        f"gemäß §§ 7, 17, 18 StVG, § 115 VVG geltend:"
    )

    for abschnitt in fliesstext.split("\n\n"):
        p = doc.add_paragraph(abschnitt.strip())
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = SCHRIFT_TEXT

    doc.add_paragraph()

    # ── 4. Schadenstabelle ────────────────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Schadenspositionen")

    positionen, gesamt_netto, gesamt_mit_haftung = _erstelle_positionen(
        schaden, haftung
    )

    if positionen:
        spalten = ["Position", "Betrag (brutto)"]
        spalten_breiten = [11.5, 3.0]
        tabelle = erstelle_positions_tabelle(
            doc, spalten, positionen, spalten_breiten
        )
        fuege_summenzeile_ein(
            tabelle,
            label="Gesamtschaden",
            wert=fmt_euro(gesamt_netto),
            n_spalten=2
        )
        if haftung < 100:
            tabelle.add_row()  # Leerzeile
            fuege_summenzeile_ein(
                tabelle,
                label=f"Forderungsbetrag ({haftung:.0f}% Haftung)",
                wert=fmt_euro(gesamt_mit_haftung),
                n_spalten=2
            )
    else:
        p = doc.add_paragraph("Schadenpositionen werden nachgereicht.")
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = GRAU

    doc.add_paragraph()

    # ── 5. Zahlungsaufforderung ───────────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Zahlungsaufforderung")

    forderungsbetrag = gesamt_mit_haftung if gesamt_mit_haftung else gesamt_netto
    zahlungstext = (
        f"Wir fordern Sie auf, den Betrag von {fmt_euro(forderungsbetrag)} "
        f"bis spätestens {frist} auf folgendes Konto zu überweisen:\n\n"
        f"Kanzlei Koch, Schatz & Kollegen\n"
        f"IBAN: DE12 3456 7890 1234 5678 90\n"
        f"BIC: COBADEFFXXX\n"
        f"Verwendungszweck: {az}"
    )

    for abschnitt in zahlungstext.split("\n\n"):
        p = doc.add_paragraph(abschnitt.strip())
        for zeile in abschnitt.strip().split("\n"):
            pass  # bereits als einzelner Paragraph
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10.5)

    doc.add_paragraph()

    # ── 6. Vorbehalt & Anlagen ────────────────────────────────────────────────
    p_vorbehalt = doc.add_paragraph(
        "Sämtliche Ansprüche werden unter ausdrücklichem Vorbehalt der "
        "Geltendmachung weiterer Schäden gestellt."
    )
    for run in p_vorbehalt.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAU

    doc.add_paragraph()

    # ── 7. Unterschriftsblock ─────────────────────────────────────────────────
    p = doc.add_paragraph("Mit freundlichen Grüßen")
    p.paragraph_format.space_after = Pt(28)
    for run in p.runs:
        run.font.size = Pt(10.5)

    p_sig = doc.add_paragraph(kanzlei["name"] if kanzlei else
                               "Rechtsanwaltskanzlei Koch, Schatz & Kollegen")
    for run in p_sig.runs:
        run.font.bold  = True
        run.font.size  = Pt(10.5)
        run.font.color.rgb = NAVY

    # ── 8. Fußzeile ───────────────────────────────────────────────────────────
    fuege_fusszeile_ein(doc, az)

    # ── Dokument in Bytes serialisieren ───────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def _vollname(beteiligter: dict) -> str:
    if not beteiligter:
        return "–"
    teile = []
    if beteiligter.get("vorname"):
        teile.append(beteiligter["vorname"])
    if beteiligter.get("name"):
        teile.append(beteiligter["name"])
    if beteiligter.get("firma") and not teile:
        return beteiligter["firma"]
    return " ".join(teile) if teile else "–"


def _erstelle_empfaenger(gegner: dict) -> list[str]:
    if not gegner:
        return ["[Versicherung des Unfallverursachers]",
                "[Straße und Hausnummer]",
                "[PLZ Ort]"]
    versicherung = gegner.get("versicherung") or "[Versicherung]"
    empf = [versicherung]
    if gegner.get("anschrift"):
        empf.append(gegner["anschrift"])
    ort = ""
    if gegner.get("plz"):
        ort += gegner["plz"] + " "
    if gegner.get("ort"):
        ort += gegner["ort"]
    if ort.strip():
        empf.append(ort.strip())
    if not ort.strip():
        empf += ["[Straße und Hausnummer]", "[PLZ Ort]"]
    return empf


def _erstelle_positionen(schaden: dict, haftung: float) -> tuple:
    """
    Erstellt Positionsliste für Schadenstabelle.
    Returns: (zeilen, gesamt_netto, gesamt_mit_haftung)
    """
    if not schaden:
        return [], None, None

    POSITIONEN_MAP = [
        ("reparaturkosten",   "Reparaturkosten (brutto, inkl. MwSt.)"),
        ("wiederbeschaffung", "Wiederbeschaffungswert"),
        ("restwert",          "abzgl. Restwert"),
        ("wertminderung",     "Merkantile Wertminderung"),
        ("nutzungsausfall",   "Nutzungsausfall"),
        ("mietwagenkosten",   "Mietwagenkosten"),
        ("sv_kosten",         "Sachverständigenkosten"),
        ("abschleppkosten",   "Abschleppkosten"),
        ("standkosten",       "Standkosten"),
        ("anabmeldekosten",   "An-/Abmeldekosten"),
        ("schmerzensgeld",    "Schmerzensgeld"),
        ("sonstiges",         schaden.get("sonstiges_beschr") or "Sonstiges"),
    ]

    zeilen = []
    gesamt = 0.0
    for feld, label in POSITIONEN_MAP:
        wert = schaden.get(feld)
        if wert is not None and wert != 0.0:
            if feld == "restwert":
                zeilen.append([label, fmt_euro(-wert)])
                gesamt -= float(wert)
            else:
                zeilen.append([label, fmt_euro(wert)])
                gesamt += float(wert)

    gesamt = round(gesamt, 2)
    gesamt_mit_haftung = round(gesamt * haftung / 100, 2) if haftung < 100 else None

    return zeilen, gesamt, gesamt_mit_haftung
