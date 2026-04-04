"""
Modul 5 – Word-Styling
========================
Wiederverwendbare Stil-Helpers für alle Word-Dokumente.
Design: Kanzlei Koch, Schatz & Kollegen (Navy #1B2A4A, Gold #C8A84B).

Alle Dokumente folgen der gleichen visuellen Sprache:
  - Briefkopf mit Kanzleiname und Kontaktdaten
  - Navy-farbige Überschriften
  - Tabellen mit Gold-Kopfzeile und alternierenden Zeilen
  - Fußzeile mit Seitenzahl und Aktenzeichen
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Farben ────────────────────────────────────────────────────────────────────

NAVY  = RGBColor(0x1B, 0x2A, 0x4A)   # Kanzlei-Navy
GOLD  = RGBColor(0xC8, 0xA8, 0x4B)   # Kanzlei-Gold
HELL  = RGBColor(0xF4, 0xF6, 0xF9)   # Helles Grau für Tabellenzeilen
WEISS = RGBColor(0xFF, 0xFF, 0xFF)
GRAU  = RGBColor(0x60, 0x60, 0x60)
SCHWARZ = RGBColor(0x1A, 0x1A, 0x1A)

# ── Seitenränder ──────────────────────────────────────────────────────────────

RAND_LINKS   = Cm(2.5)
RAND_RECHTS  = Cm(2.0)
RAND_OBEN    = Cm(2.5)
RAND_UNTEN   = Cm(2.0)

# ── Schriften ─────────────────────────────────────────────────────────────────

SCHRIFT_TEXT   = "Calibri"
SCHRIFT_KOPF   = "Calibri"


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def _rgb_hex(farbe: RGBColor) -> str:
    return f"{farbe[0]:02X}{farbe[1]:02X}{farbe[2]:02X}"


def setze_zellen_farbe(zelle, hex_farbe: str):
    """Setzt die Hintergrundfarbe einer Tabellenzelle."""
    tc = zelle._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_farbe)
    tcPr.append(shd)


def setze_zellen_rand(zelle, seiten=("top","bottom","left","right"),
                       groesse=4, farbe="1B2A4A"):
    """Setzt Ränder einer Tabellenzelle."""
    tc = zelle._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for seite in seiten:
        border = OxmlElement(f"w:{seite}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    str(groesse))
        border.set(qn("w:color"), farbe)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def entferne_tabellen_raender(tabelle):
    """Entfernt alle äußeren Ränder einer Tabelle."""
    tbl = tabelle._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for seite in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{seite}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ── Dokument-Grundkonfiguration ───────────────────────────────────────────────

def erstelle_dokument() -> Document:
    """Erstellt ein neues Dokument mit Kanzlei-Seitenrändern."""
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.left_margin   = RAND_LINKS
        section.right_margin  = RAND_RECHTS
        section.top_margin    = RAND_OBEN
        section.bottom_margin = RAND_UNTEN

    # Standard-Schrift für Normal-Style
    style = doc.styles["Normal"]
    style.font.name = SCHRIFT_TEXT
    style.font.size = Pt(10.5)
    style.font.color.rgb = SCHWARZ

    return doc


# ── Briefkopf ─────────────────────────────────────────────────────────────────

def fuege_briefkopf_ein(doc: Document, kanzlei_info: dict = None):
    """
    Fügt einen professionellen Kanzlei-Briefkopf ein.

    kanzlei_info: Dict mit name, strasse, ort, telefon, email, iban
    """
    if kanzlei_info is None:
        kanzlei_info = {
            "name":     "Rechtsanwaltskanzlei Koch, Schatz & Kollegen",
            "strasse":  "Frankfurter Straße 12",
            "ort":      "63065 Offenbach am Main",
            "telefon":  "069 / 83 10 99 - 0",
            "fax":      "069 / 83 10 99 - 99",
            "email":    "info@anwalt-offenbach.de",
            "web":      "www.anwalt-offenbach.de",
        }

    # Kopf-Tabelle: Kanzleiname links, Kontakt rechts
    tabelle = doc.add_table(rows=1, cols=2)
    tabelle.style = "Table Grid"
    entferne_tabellen_raender(tabelle)

    # Linke Spalte: Kanzleiname groß
    links = tabelle.rows[0].cells[0]
    p = links.paragraphs[0]
    run = p.add_run(kanzlei_info["name"])
    run.font.name   = SCHRIFT_KOPF
    run.font.size   = Pt(14)
    run.font.bold   = True
    run.font.color.rgb = NAVY

    # Rechte Spalte: Kontaktdaten klein
    rechts = tabelle.rows[0].cells[1]
    rechts.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for zeile in [
        kanzlei_info.get("strasse", ""),
        kanzlei_info.get("ort", ""),
        f'Tel.: {kanzlei_info.get("telefon", "")}',
        f'E-Mail: {kanzlei_info.get("email", "")}',
    ]:
        p = rechts.add_paragraph(zeile)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = GRAU

    # Trennlinie in Navy
    p_linie = doc.add_paragraph()
    pPr = p_linie._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:color"), _rgb_hex(NAVY))
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()  # Abstand nach Kopf


# ── Absenderblock (Anwalt) & Empfängerblock ───────────────────────────────────

def fuege_adressblock_ein(doc: Document, empfaenger: list[str],
                           absender_zusatz: list[str] = None,
                           betreff: str = None,
                           aktenzeichen: str = None,
                           datum: str = None):
    """
    Fügt Absender-/Empfänger-Adressblock und Betreff ein.
    Standard Brieflayout (DIN 5008).
    """
    # Datum und Aktenzeichen (rechtsbündig)
    if datum or aktenzeichen:
        tabelle = doc.add_table(rows=1, cols=2)
        tabelle.style = "Table Grid"
        entferne_tabellen_raender(tabelle)
        links = tabelle.rows[0].cells[0]
        rechts = tabelle.rows[0].cells[1]
        rechts.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if aktenzeichen:
            p = rechts.paragraphs[0]
            run = p.add_run(f"Az.: {aktenzeichen}")
            run.font.size = Pt(9)
            run.font.color.rgb = GRAU
        if datum:
            p2 = rechts.add_paragraph(datum)
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p2.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = GRAU

    # Empfänger
    doc.add_paragraph()
    for zeile in empfaenger:
        p = doc.add_paragraph(zeile)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10.5)

    doc.add_paragraph()

    # Betreff (fett, Navy)
    if betreff:
        p = doc.add_paragraph()
        run = p.add_run(betreff)
        run.font.bold  = True
        run.font.size  = Pt(11)
        run.font.color.rgb = NAVY
        doc.add_paragraph()


# ── Überschriften ─────────────────────────────────────────────────────────────

def fuege_abschnittstitel_ein(doc: Document, titel: str):
    """Fügt eine Navy-formatierte Abschnittsüberschrift ein."""
    p = doc.add_paragraph()
    run = p.add_run(titel)
    run.font.name  = SCHRIFT_KOPF
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


# ── Tabellen ──────────────────────────────────────────────────────────────────

def erstelle_positions_tabelle(doc: Document, spalten: list[str],
                                zeilen: list[list],
                                spalten_breiten: list[float] = None) -> object:
    """
    Erstellt eine formatierte Positionstabelle.

    spalten: Spaltenköpfe
    zeilen:  Liste von Zeilen (jede Zeile = Liste von Werten)
    spalten_breiten: Breite in cm (optional)

    Returns: docx-Tabelle
    """
    tabelle = doc.add_table(rows=1 + len(zeilen), cols=len(spalten))
    tabelle.style = "Table Grid"
    entferne_tabellen_raender(tabelle)

    # Spaltenbreiten setzen
    if spalten_breiten:
        for i, breite in enumerate(spalten_breiten):
            for zeile_obj in tabelle.rows:
                zeile_obj.cells[i].width = Cm(breite)

    # Kopfzeile (Gold-Hintergrund, weiße Schrift)
    kopf_zeile = tabelle.rows[0]
    for i, kopf in enumerate(spalten):
        zelle = kopf_zeile.cells[i]
        setze_zellen_farbe(zelle, _rgb_hex(NAVY))
        p = zelle.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(kopf)
        run.font.bold  = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = WEISS
        run.font.name  = SCHRIFT_KOPF

    # Datenzeilen
    for z_idx, zeile_daten in enumerate(zeilen):
        tabellen_zeile = tabelle.rows[z_idx + 1]
        farbe_hex = _rgb_hex(HELL) if z_idx % 2 == 0 else "FFFFFF"
        for s_idx, wert in enumerate(zeile_daten):
            zelle = tabellen_zeile.cells[s_idx]
            setze_zellen_farbe(zelle, farbe_hex)
            p = zelle.paragraphs[0]
            # Betragspalten rechtsbündig
            if isinstance(wert, float) or (
                isinstance(wert, str) and re.match(r"[\d.,]+\s*€?$", wert.strip())
            ):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(wert))
            run.font.size = Pt(9.5)
            run.font.name = SCHRIFT_TEXT

    return tabelle


def fuege_summenzeile_ein(tabelle, label: str, wert: str,
                           n_spalten: int):
    """Fügt eine fett formatierte Summenzeile an eine Tabelle an."""
    zeile = tabelle.add_row()
    # Erste Zellen leer (bis auf letzte zwei)
    for i in range(n_spalten - 2):
        setze_zellen_farbe(zeile.cells[i], "FFFFFF")

    # Label
    lbl_zelle = zeile.cells[n_spalten - 2]
    setze_zellen_farbe(lbl_zelle, _rgb_hex(NAVY))
    p = lbl_zelle.paragraphs[0]
    run = p.add_run(label)
    run.font.bold  = True
    run.font.size  = Pt(9.5)
    run.font.color.rgb = WEISS

    # Wert
    wert_zelle = zeile.cells[n_spalten - 1]
    setze_zellen_farbe(wert_zelle, _rgb_hex(NAVY))
    p2 = wert_zelle.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(wert)
    run2.font.bold  = True
    run2.font.size  = Pt(9.5)
    run2.font.color.rgb = WEISS


# ── Fußzeile ──────────────────────────────────────────────────────────────────

def fuege_fusszeile_ein(doc: Document, aktenzeichen: str = None):
    """Fügt eine Fußzeile mit Seitenzahl ein."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Trennlinie oben
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "6")
    top.set(qn("w:color"), _rgb_hex(NAVY))
    pBdr.append(top)
    pPr.append(pBdr)

    # Text: "Az.: 42/25  |  Seite X"
    if aktenzeichen:
        run_az = p.add_run(f"Az.: {aktenzeichen}    |    ")
        run_az.font.size = Pt(8)
        run_az.font.color.rgb = GRAU

    run_seite = p.add_run("Seite ")
    run_seite.font.size = Pt(8)
    run_seite.font.color.rgb = GRAU

    # Automatische Seitenzahl-Felder
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_feld = p.add_run()
    run_feld._r.append(fldChar1)
    run_feld._r.append(instrText)
    run_feld._r.append(fldChar2)
    run_feld.font.size = Pt(8)
    run_feld.font.color.rgb = GRAU


# ── Hilfsfunktion: Euro-Formatierung ─────────────────────────────────────────

def fmt_euro(betrag) -> str:
    """Formatiert einen Betrag als deutsches Euro-Format: 1.234,56 €"""
    if betrag is None:
        return "–"
    if isinstance(betrag, (int, float)):
        return f"{betrag:,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")
    return str(betrag)


def fmt_datum(datum_str: str) -> str:
    """Konvertiert ISO-Datum (2025-03-15) in deutsches Format (15.03.2025)."""
    if not datum_str:
        return "–"
    teile = datum_str.split("-")
    if len(teile) == 3:
        return f"{teile[2]}.{teile[1]}.{teile[0]}"
    return datum_str
