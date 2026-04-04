"""
Modul 5 – Generator: Sachstandsanfrage
========================================
Erzeugt eine Sachstandsanfrage an den Versicherer.
Wird genutzt wenn keine Reaktion auf das Forderungsschreiben erfolgt.

Struktur:
  1. Briefkopf
  2. Adressblock + Betreff
  3. Bezugnahme auf Forderungsschreiben
  4. Aktueller Regulierungsstand (falls vorhanden)
  5. Fristsetzung
  6. Androhung gerichtlicher Schritte
  7. Unterschriftsblock + Fußzeile
"""

import io
import logging
from datetime import date, timedelta
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styling import (
    erstelle_dokument, fuege_briefkopf_ein, fuege_adressblock_ein,
    fuege_abschnittstitel_ein, erstelle_positions_tabelle,
    fuege_fusszeile_ein, fmt_euro, fmt_datum,
    NAVY, GRAU, SCHWARZ, SCHRIFT_TEXT
)

logger = logging.getLogger(__name__)


def generiere_sachstandsanfrage(akte_daten: dict) -> bytes:
    """
    Erstellt eine Sachstandsanfrage als Word-Dokument.

    Args:
        akte_daten: Dict mit akte, mandant, gegner, schaden,
                    regulierungen, kanzlei.
                    Optional: brieftext (str) – überschreibt den generierten
                    Briefkörper mit dem angegebenen Text (PRD-25d).

    Returns:
        Word-Dokument als Bytes
    """
    doc = erstelle_dokument()
    akte          = akte_daten.get("akte", {})
    mandant       = akte_daten.get("mandant")
    gegner        = akte_daten.get("gegner")
    schaden       = akte_daten.get("schaden")
    regulierungen = akte_daten.get("regulierungen", [])
    kanzlei       = akte_daten.get("kanzlei")
    brieftext     = akte_daten.get("brieftext")

    az          = akte.get("aktenzeichen", "–")
    unfalldatum = fmt_datum(akte.get("unfalldatum", ""))
    unfallort   = akte.get("unfallort", "")
    heute       = date.today().strftime("%d.%m.%Y")
    frist       = (date.today() + timedelta(days=7)).strftime("%d.%m.%Y")

    # Versicherungsnummer des Gegners
    schaden_nr = (gegner or {}).get("schaden_nr", "–")
    versicherung = (gegner or {}).get("versicherung", "[Versicherung]")

    # ── Briefkopf ─────────────────────────────────────────────────────────────
    fuege_briefkopf_ein(doc, kanzlei)

    # ── Empfänger ─────────────────────────────────────────────────────────────
    empf = _empfaenger(gegner)
    betreff = (
        f"Sachstandsanfrage\n"
        f"Verkehrsunfall vom {unfalldatum}"
        + (f" in {unfallort}" if unfallort else "")
        + f"\nIhr Zeichen: {schaden_nr} | Unser Zeichen: {az}"
    )
    fuege_adressblock_ein(
        doc, empf, betreff=betreff, aktenzeichen=az, datum=heute
    )

    # ── Briefkörper ───────────────────────────────────────────────────────────
    # PRD-25d: Wenn brieftext übergeben, diesen als Briefkörper verwenden
    if brieftext:
        for absatz in brieftext.split("\n\n"):
            absatz = absatz.strip()
            if absatz:
                p = doc.add_paragraph(absatz)
                _stil(p)
                doc.add_paragraph()

        # Direkt zur Unterschrift springen
        p = doc.add_paragraph("Mit freundlichen Grüßen")
        p.paragraph_format.space_after = Pt(28)
        _stil(p)
        p2 = doc.add_paragraph(
            kanzlei["name"] if kanzlei else
            "Rechtsanwaltskanzlei Koch, Schatz & Kollegen"
        )
        for run in p2.runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
        fuege_fusszeile_ein(doc, az)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Eröffnung (Standard-Pfad) ─────────────────────────────────────────────
    mandant_name = _vollname(mandant) if mandant else "unser Mandant"

    p = doc.add_paragraph(
        f"in vorbezeichneter Angelegenheit haben wir Ihnen bereits mit "
        f"Schreiben unsere Schadensersatzansprüche für {mandant_name} "
        f"angezeigt."
    )
    _stil(p)
    doc.add_paragraph()

    # ── Aktueller Stand ───────────────────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Aktueller Schadensstand")

    gesamtschaden = (schaden or {}).get("gesamt_brutto")
    if gesamtschaden:
        p = doc.add_paragraph(
            f"Der von uns geltend gemachte Gesamtschaden beläuft sich auf "
            f"{fmt_euro(gesamtschaden)}."
        )
        _stil(p)

    if regulierungen:
        # Regulierungstabelle
        doc.add_paragraph()
        zeilen = []
        for reg in regulierungen:
            zeilen.append([
                fmt_datum(reg.get("datum", "")),
                fmt_euro(reg.get("betrag_gefordert")),
                fmt_euro(reg.get("betrag_reguliert")),
                reg.get("status", "–").replace("_", " ").capitalize(),
            ])
        erstelle_positions_tabelle(
            doc,
            spalten=["Datum", "Gefordert", "Reguliert", "Status"],
            zeilen=zeilen,
            spalten_breiten=[3.5, 3.5, 3.5, 3.5],
        )

        gesamt_reguliert = sum(
            float(r.get("betrag_reguliert") or 0) for r in regulierungen
        )
        gesamt_gefordert = sum(
            float(r.get("betrag_gefordert") or 0) for r in regulierungen
        )
        differenz = gesamt_gefordert - gesamt_reguliert

        doc.add_paragraph()
        if differenz > 0.01:
            p = doc.add_paragraph(
                f"Noch ausstehend: {fmt_euro(round(differenz, 2))}"
            )
            _stil(p, fett=True)
    else:
        p = doc.add_paragraph(
            "Bisher ist von Ihnen weder eine Zahlung noch eine inhaltliche "
            "Stellungnahme eingegangen."
        )
        _stil(p)

    doc.add_paragraph()

    # ── Fristsetzung ──────────────────────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Fristsetzung")

    p = doc.add_paragraph(
        f"Wir setzen Ihnen hiermit eine Nachfrist bis zum {frist} "
        f"und fordern Sie auf, den noch offenen Betrag bis zu diesem Datum "
        f"auf unser in unserem Forderungsschreiben genanntes Konto zu überweisen "
        f"bzw. uns Ihre Regulierungsbereitschaft schriftlich mitzuteilen."
    )
    _stil(p)

    doc.add_paragraph()

    # ── Androhung Klage ───────────────────────────────────────────────────────
    p = doc.add_paragraph(
        f"Wir weisen darauf hin, dass wir nach fruchtlosem Ablauf der "
        f"gesetzten Frist ohne weitere Ankündigung gerichtliche Schritte "
        f"einleiten werden. Die anfallenden Gerichts- und Anwaltskosten "
        f"gehen in diesem Fall zu Ihren Lasten."
    )
    _stil(p)

    doc.add_paragraph()

    # ── Unterschrift ──────────────────────────────────────────────────────────
    p = doc.add_paragraph("Mit freundlichen Grüßen")
    p.paragraph_format.space_after = Pt(28)
    _stil(p)

    p2 = doc.add_paragraph(
        kanzlei["name"] if kanzlei else
        "Rechtsanwaltskanzlei Koch, Schatz & Kollegen"
    )
    for run in p2.runs:
        run.font.bold = True
        run.font.color.rgb = NAVY

    fuege_fusszeile_ein(doc, az)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def _vollname(b: dict) -> str:
    teile = [b.get("vorname", ""), b.get("name", "")]
    name = " ".join(t for t in teile if t).strip()
    return name or b.get("firma", "–") or "–"


def _empfaenger(gegner: dict) -> list[str]:
    if not gegner:
        return ["[Versicherung]", "[PLZ Ort]"]
    zeilen = [gegner.get("versicherung") or "[Versicherung]"]
    if gegner.get("anschrift"):
        zeilen.append(gegner["anschrift"])
    ort = " ".join(filter(None, [gegner.get("plz"), gegner.get("ort")]))
    if ort:
        zeilen.append(ort)
    return zeilen


def _stil(p, fett: bool = False):
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = SCHRIFT_TEXT
        if fett:
            run.font.bold = True
