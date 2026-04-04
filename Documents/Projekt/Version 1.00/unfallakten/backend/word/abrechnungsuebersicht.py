"""
Modul 5 – Generator: Abrechnungsübersicht
===========================================
Erzeugt eine vollständige Abrechnungsübersicht für den Mandanten.
Zeigt alle Schadenpositionen, Regulierungen und den ausstehenden Betrag.

Struktur:
  1. Briefkopf
  2. Adressblock (Mandant)
  3. Betreff
  4. Einleitung
  5. Schadenstabelle (alle Positionen)
  6. Regulierungsverlauf (Tabelle)
  7. Zusammenfassung (gefordert / reguliert / ausstehend)
  8. Hinweis auf weitere Schritte
  9. Unterschrift + Fußzeile

v8-Änderungen:
  - Schadenfelder um rep_gutachten_netto, rep_rechnung_netto, verdienstausfall,
    haushalt, unkostenpauschale, kostennb erweitert
  - Haftungsquote wird auf Gesamtschaden angewendet
  - Regulierungen kommen aus akte_daten["abrechnungen"] (v8) mit Fallback
    auf akte_daten["regulierungen"] (alt)
"""

import io
import logging
from datetime import date
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styling import (
    erstelle_dokument, fuege_briefkopf_ein, fuege_adressblock_ein,
    fuege_abschnittstitel_ein, erstelle_positions_tabelle,
    fuege_summenzeile_ein, fuege_fusszeile_ein,
    fmt_euro, fmt_datum, NAVY, GOLD, GRAU, SCHRIFT_TEXT,
    setze_zellen_farbe, _rgb_hex
)

logger = logging.getLogger(__name__)

# Ampelfarben für Status
GRUEN  = RGBColor(0x1A, 0x7A, 0x3C)
ROT    = RGBColor(0xC0, 0x30, 0x2A)
ORANGE = RGBColor(0xD4, 0x7E, 0x10)


def generiere_abrechnungsuebersicht(akte_daten: dict) -> bytes:
    """
    Erstellt eine Abrechnungsübersicht als Word-Dokument.

    Args:
        akte_daten: Dict mit akte, mandant, gegner, schaden,
                    abrechnungen (v8) | regulierungen (alt), kanzlei

    Returns:
        Word-Dokument als Bytes
    """
    doc = erstelle_dokument()
    akte      = akte_daten.get("akte", {})
    mandant   = akte_daten.get("mandant")
    schaden   = akte_daten.get("schaden")
    kanzlei   = akte_daten.get("kanzlei")

    # v8: abrechnungen bevorzugen, Fallback auf altes regulierungen-Format
    regulierungen = akte_daten.get("abrechnungen") or akte_daten.get("regulierungen", [])

    az          = akte.get("aktenzeichen", "–")
    unfalldatum = fmt_datum(akte.get("unfalldatum", ""))
    unfallort   = akte.get("unfallort", "")
    status      = akte.get("status", "offen")
    hq          = float(akte.get("haftungsquote") or 100) / 100
    heute       = date.today().strftime("%d.%m.%Y")

    # ── Briefkopf ─────────────────────────────────────────────────────────────
    fuege_briefkopf_ein(doc, kanzlei)

    # ── Empfänger: Mandant ────────────────────────────────────────────────────
    empf = _empfaenger_mandant(mandant)
    mandant_name = _vollname(mandant) if mandant else "Mandant"

    betreff = (
        f"Abrechnungsübersicht – Ihr Unfallschaden\n"
        f"Verkehrsunfall vom {unfalldatum}"
        + (f" in {unfallort}" if unfallort else "")
        + f"\nUnser Aktenzeichen: {az}"
    )
    fuege_adressblock_ein(
        doc, empf, betreff=betreff, aktenzeichen=az, datum=heute
    )

    # ── Einleitung ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        f"sehr geehrte/r {mandant_name},\n\n"
        f"wir übersenden Ihnen hiermit eine aktuelle Übersicht über den "
        f"Stand der Regulierung Ihres Unfallschadens vom {unfalldatum}."
    )
    _stil(p)
    doc.add_paragraph()

    # ── Status-Badge ──────────────────────────────────────────────────────────
    _fuege_status_badge_ein(doc, status)
    doc.add_paragraph()

    # ── Schadenstabelle ───────────────────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Ihre Schadenpositionen")

    gesamt_brutto = 0.0
    if schaden:
        positionen, gesamt_brutto = _erstelle_schaden_zeilen(schaden)
        if positionen:
            tabelle = erstelle_positions_tabelle(
                doc,
                spalten=["Schadenposition", "Betrag (brutto)"],
                zeilen=positionen,
                spalten_breiten=[11.5, 3.0],
            )
            fuege_summenzeile_ein(
                tabelle,
                label="Gesamtschaden (brutto)",
                wert=fmt_euro(gesamt_brutto),
                n_spalten=2,
            )
    else:
        p = doc.add_paragraph("Schadenpositionen sind noch nicht vollständig erfasst.")
        _stil(p, grau=True)

    # Haftungsquote anwenden
    gesamt_netto = round(gesamt_brutto * hq, 2)

    # Haftungsquote-Hinweis wenn nicht 100 %
    if hq < 1.0:
        doc.add_paragraph()
        p_hq = doc.add_paragraph(
            f"Nach der vereinbarten Haftungsquote von {int(round(hq * 100))} % "
            f"ergibt sich ein geltend gemachter Schaden von {fmt_euro(gesamt_netto)}."
        )
        _stil(p_hq)

    doc.add_paragraph()

    # ── Regulierungsverlauf ───────────────────────────────────────────────────
    if regulierungen:
        fuege_abschnittstitel_ein(doc, "Regulierungsverlauf")

        reg_zeilen = []
        for reg in regulierungen:
            # v8-Felder (abrechnungsschreiben) und altes Format (regulierungen)
            gef = float(reg.get("betrag_gefordert") or reg.get("gesamt_gefordert") or 0)
            reg_b = float(reg.get("betrag_reguliert") or reg.get("gesamt_reguliert") or 0)
            differenz = round(gef - reg_b, 2)
            reg_status = reg.get("status", "–")
            # Versicherung/Quelle als Zusatzinfo falls vorhanden
            vers = reg.get("versicherung", "") or ""
            datum_str = fmt_datum(reg.get("datum", ""))
            if vers:
                datum_str = f"{datum_str}\n{vers}"
            reg_zeilen.append([
                datum_str,
                fmt_euro(gef),
                fmt_euro(reg_b),
                fmt_euro(differenz),
                _status_label(reg_status),
            ])

        erstelle_positions_tabelle(
            doc,
            spalten=["Datum / Versicherung", "Gefordert", "Reguliert", "Differenz", "Status"],
            zeilen=reg_zeilen,
            spalten_breiten=[3.5, 2.8, 2.8, 2.4, 3.0],
        )
        doc.add_paragraph()

    # ── Zusammenfassung ───────────────────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Zusammenfassung")

    gesamt_reguliert = sum(
        float(r.get("betrag_reguliert") or r.get("gesamt_reguliert") or 0)
        for r in regulierungen
    )
    ausstehend = round(gesamt_netto - gesamt_reguliert, 2)

    zusammen_zeilen = [
        ["Gesamtschaden brutto", fmt_euro(round(gesamt_brutto, 2))],
    ]
    if hq < 1.0:
        zusammen_zeilen.append([
            f"Geltend gemacht ({int(round(hq * 100))} % Haftungsquote)",
            fmt_euro(gesamt_netto),
        ])
    zusammen_zeilen.append(["Bereits reguliert", fmt_euro(round(gesamt_reguliert, 2))])

    tabelle = erstelle_positions_tabelle(
        doc,
        spalten=["Position", "Betrag"],
        zeilen=zusammen_zeilen,
        spalten_breiten=[11.5, 3.0],
    )

    # Ausstehend-Zeile farblich hervorheben
    farbe_hex = _rgb_hex(GRUEN) if ausstehend <= 0 else _rgb_hex(ROT)
    zeile_obj = tabelle.add_row()
    setze_zellen_farbe(zeile_obj.cells[0], farbe_hex)
    setze_zellen_farbe(zeile_obj.cells[1], farbe_hex)

    p_lbl = zeile_obj.cells[0].paragraphs[0]
    run_lbl = p_lbl.add_run(
        "Noch ausstehend" if ausstehend > 0 else "Vollständig reguliert ✓"
    )
    run_lbl.font.bold = True
    run_lbl.font.size = Pt(9.5)
    run_lbl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p_wert = zeile_obj.cells[1].paragraphs[0]
    p_wert.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_wert = p_wert.add_run(fmt_euro(abs(round(ausstehend, 2))))
    run_wert.font.bold = True
    run_wert.font.size = Pt(9.5)
    run_wert.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # ── Hinweise / nächste Schritte ───────────────────────────────────────────
    fuege_abschnittstitel_ein(doc, "Weiteres Vorgehen")

    if ausstehend > 0.01:
        text = (
            f"Der Versicherer hat den Schaden bislang nicht vollständig reguliert. "
            f"Wir werden die noch ausstehenden {fmt_euro(ausstehend)} weiter "
            f"einfordern. Sie werden von uns informiert, sobald sich der Sachstand "
            f"ändert."
        )
    else:
        text = (
            f"Ihr Unfallschaden wurde vollständig reguliert. "
            f"Wir haben die Akte als abgeschlossen vermerkt. "
            f"Bei Rückfragen stehen wir Ihnen jederzeit zur Verfügung."
        )

    p = doc.add_paragraph(text)
    _stil(p)

    doc.add_paragraph()

    # ── Unterschrift ──────────────────────────────────────────────────────────
    p = doc.add_paragraph("Mit freundlichen Grüßen")
    p.paragraph_format.space_after = Pt(28)
    _stil(p)

    p2 = doc.add_paragraph(
        kanzlei["name"] if kanzlei else
        "Rechtsanwaltskanzlei Koch, Schatz & Kollegen"
    )
    for run in p2.runs:
        run.font.bold = True
        run.font.color.rgb = NAVY

    fuege_fusszeile_ein(doc, az)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def _vollname(b: dict) -> str:
    teile = [b.get("vorname", ""), b.get("name", "")]
    name = " ".join(t for t in teile if t).strip()
    return name or b.get("firma", "–") or "–"


def _empfaenger_mandant(mandant: dict) -> list[str]:
    if not mandant:
        return ["[Mandant]"]
    zeilen = []
    name = _vollname(mandant)
    if name and name != "–":
        zeilen.append(name)
    if mandant.get("anschrift"):
        zeilen.append(mandant["anschrift"])
    ort = " ".join(filter(None, [mandant.get("plz"), mandant.get("ort")]))
    if ort:
        zeilen.append(ort)
    return zeilen or ["[Mandant]"]


def _erstelle_schaden_zeilen(schaden: dict) -> tuple:
    """
    Erzeugt Tabellenzeilen für alle Schadenpositionen.
    v8: rep_gutachten_netto / rep_rechnung_netto haben Vorrang vor reparaturkosten.
    Neue Felder: verdienstausfall, haushalt, unkostenpauschale, kostennb.
    """
    # Fahrzeugschaden: rep-Felder bevorzugen
    rep_gutachten = float(schaden.get("rep_gutachten_netto") or 0)
    rep_rechnung  = float(schaden.get("rep_rechnung_netto")  or 0)
    rep_alt       = float(schaden.get("reparaturkosten")     or 0)

    zeilen = []
    gesamt = 0.0

    # ── Reparatur / Fahrzeugschaden ───────────────────────────────────────────
    if rep_rechnung > 0:
        zeilen.append(["Reparaturkosten lt. Rechnung (netto)", fmt_euro(rep_rechnung)])
        gesamt += rep_rechnung
    elif rep_gutachten > 0:
        zeilen.append(["Reparaturkosten lt. Gutachten (netto)", fmt_euro(rep_gutachten)])
        gesamt += rep_gutachten
    elif rep_alt > 0:
        zeilen.append(["Reparaturkosten (inkl. MwSt.)", fmt_euro(rep_alt)])
        gesamt += rep_alt

    # ── Restliche Felder in fester Reihenfolge ────────────────────────────────
    WEITERE = [
        ("wiederbeschaffung", "Wiederbeschaffungswert"),
        ("restwert",          "abzgl. Restwert"),          # wird als Minus gerechnet
        ("wertminderung",     "Merkantile Wertminderung"),
        ("nutzungsausfall",   "Nutzungsausfall"),
        ("mietwagenkosten",   "Mietwagenkosten"),
        ("sv_kosten",         "Sachverständigenkosten"),
        ("abschleppkosten",   "Abschleppkosten"),
        ("standkosten",       "Standkosten"),
        ("anabmeldekosten",   "An-/Abmeldekosten"),
        ("verdienstausfall",  "Verdienstausfall"),
        ("haushalt",          "Haushaltsführungsschaden"),
        ("schmerzensgeld",    "Schmerzensgeld"),
        ("unkostenpauschale", "Unkostenpauschale"),
        ("kostennb",          "Kostennebenforderungen"),
        ("sonstiges",         schaden.get("sonstiges_beschr") or "Sonstiges"),
    ]

    for feld, label in WEITERE:
        wert = schaden.get(feld)
        if wert is None or float(wert) == 0.0:
            continue
        wert_f = float(wert)
        if feld == "restwert":
            zeilen.append([label, fmt_euro(-wert_f)])
            gesamt -= wert_f
        else:
            zeilen.append([label, fmt_euro(wert_f)])
            gesamt += wert_f

    return zeilen, round(gesamt, 2)


def _fuege_status_badge_ein(doc, status: str):
    """Fügt einen farbigen Status-Hinweis-Block ein."""
    STATUS_MAP = {
        "offen":           ("Akte offen – Regulierung läuft",      ORANGE),
        "in_regulierung":  ("Regulierung in Bearbeitung",           ORANGE),
        "abgeschlossen":   ("Regulierung abgeschlossen ✓",          GRUEN),
        "klage":           ("Klage eingereicht",                    ROT),
    }
    text, farbe = STATUS_MAP.get(status, ("Status unbekannt", GRAU))

    tabelle = doc.add_table(rows=1, cols=1)
    tabelle.style = "Table Grid"
    from .styling import entferne_tabellen_raender
    entferne_tabellen_raender(tabelle)
    zelle = tabelle.rows[0].cells[0]
    setze_zellen_farbe(zelle, _rgb_hex(farbe))
    p = zelle.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"  {text}  ")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _status_label(status: str) -> str:
    return {
        "teilreguliert":  "Teilreguliert",
        "vollreguliert":  "Vollreguliert ✓",
        "abgelehnt":      "Abgelehnt",
        "ausstehend":     "Ausstehend",
        "erfolgreich":    "Reguliert",      # parse_status aus v8
    }.get(status, status.capitalize() if status else "–")


def _stil(p, grau: bool = False):
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = SCHRIFT_TEXT
        if grau:
            run.font.color.rgb = GRAU
