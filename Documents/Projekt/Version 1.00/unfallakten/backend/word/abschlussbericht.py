"""
Abschluss-/Sachstandsbericht – DOCX-Renderer
=============================================
Rendert das Übersichts-Objekt (services/abschluss_uebersicht.py) im
Kanzlei-Hausstil (styling.py, wie Sachstandsanfrage). Der Renderer ist
"dumm": keine eigene Rechenlogik.

Anatomie (Spec §9): Briefkopf/Betreff → Ergebnis bzw. Arbeitsstand →
"Was bei Ihnen ankommt" (nur Abschluss) → Gegenüberstellung + Zahlungs-
verlauf → Anwaltskosten → Schluss (+ Bewertungszeile) → Grußformel.
"""
import io
from datetime import date

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..services.abschluss_uebersicht import baue_abschluss_uebersicht
from .styling import (
    erstelle_dokument, fuege_briefkopf_ein, fuege_adressblock_ein,
    fuege_abschnittstitel_ein, erstelle_positions_tabelle,
    fuege_fusszeile_ein, setze_zellen_farbe, fmt_euro, fmt_datum,
    NAVY, GRAU,
)

_GOLD_HELL = "F7F1DF"

_STATUS_LABEL = {
    "voll":     "vollständig gezahlt",
    "gekuerzt": "gekürzt",
    "offen":    "noch offen",
    "abzug":    "Abzugsposten",
}


def dateiendung() -> str:
    return "docx"


def _anrede_zeile(mandant: dict) -> str:
    anrede = (mandant.get("anrede") or "").strip()
    name = mandant.get("name") or ""
    nachname = name.split()[-1] if name else ""
    if anrede == "1" and nachname:
        return f"Sehr geehrter Herr {nachname},"
    if anrede == "2" and nachname:
        return f"Sehr geehrte Frau {nachname},"
    return "Sehr geehrte Damen und Herren,"


def _absatz(doc, text, size=10.5, bold=False, farbe=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if farbe is not None:
        run.font.color.rgb = farbe
    return p


def _ergebnis_kachel(doc, zeilen):
    tab = doc.add_table(rows=1, cols=1)
    tab.style = "Table Grid"
    zelle = tab.rows[0].cells[0]
    setze_zellen_farbe(zelle, _GOLD_HELL)
    for i, (text, gross) in enumerate(zeilen):
        p = zelle.paragraphs[0] if i == 0 else zelle.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = gross
        run.font.size = Pt(13 if gross else 10)
        run.font.color.rgb = NAVY
    doc.add_paragraph()


def generiere_abschlussbericht(akte_daten: dict) -> bytes:
    ueb = baue_abschluss_uebersicht(akte_daten)
    modus = ueb["modus"]
    az = ueb["akte"]["az"]
    summen = ueb["summen"]

    doc = erstelle_dokument()
    fuege_briefkopf_ein(doc, akte_daten.get("kanzlei"))

    mandant = ueb["mandant"]
    empfaenger = [z for z in (mandant["name"], mandant["anschrift"],
                              mandant["plz_ort"]) if z]
    betreff = (f"Abschluss Ihrer Schadenersatzangelegenheit — "
               f"Unfall vom {fmt_datum(ueb['akte']['unfalltag'])}"
               if modus == "abschluss" else
               f"Sachstandsbericht zu Ihrer Schadenersatzangelegenheit — "
               f"Unfall vom {fmt_datum(ueb['akte']['unfalltag'])}")
    fuege_adressblock_ein(
        doc, empfaenger, betreff=betreff, aktenzeichen=az,
        datum=date.today().strftime("%d.%m.%Y"))

    _absatz(doc, _anrede_zeile(mandant))
    doc.add_paragraph()

    if modus == "abschluss":
        _ergebnis_kachel(doc, [
            (f"Für Sie durchgesetzt: {fmt_euro(summen['gezahlt'])}", True),
            (f"von {fmt_euro(summen['gefordert'])} geforderten "
             f"Schadenersatzansprüchen", False),
        ])
        fuege_abschnittstitel_ein(doc, "Was davon bei Ihnen ankommt")
        _absatz(doc, f"Insgesamt reguliert wurden "
                     f"{fmt_euro(summen['gezahlt'])} — davon gingen "
                     f"{fmt_euro(summen['an_mandant'])} direkt an Sie.")
        if summen["an_dritte"] > 0.005:
            _absatz(doc, f"Die übrigen {fmt_euro(summen['an_dritte'])} wurden "
                         f"unmittelbar an Dritte gezahlt (z. B. Werkstatt, "
                         f"Sachverständiger, Mietwagenunternehmen).")
    else:
        fuege_abschnittstitel_ein(doc, "Woran wir arbeiten / worauf wir warten")
        offene = [p for p in ueb["positionen"] if p["status"] == "offen"]
        erledigte = [p for p in ueb["positionen"] if p["status"] == "voll"]
        for pos in erledigte:
            _absatz(doc, f"✓ {pos['label']} — erledigt", size=10)
        for pos in offene:
            _absatz(doc, f"○ {pos['label']} — noch offen "
                         f"({fmt_euro(pos['gefordert'])})", size=10)
        if ueb["schluss"]["naechste_schritte_text"]:
            _absatz(doc, f"Nächster Schritt: "
                         f"{ueb['schluss']['naechste_schritte_text']}", bold=True)

    fuege_abschnittstitel_ein(doc, "Gegenüberstellung Ihrer Ansprüche")
    zeilen = []
    for p in ueb["positionen"]:
        grund = p["kuerzung_grund"] or _STATUS_LABEL[p["status"]]
        zeilen.append([
            p["label"],
            fmt_euro(p["gefordert"]),
            fmt_euro(p["gezahlt"]) if p["gezahlt"] is not None else "–",
            fmt_euro(p["differenz"]) if p["differenz"] > 0.005 else "–",
            grund,
        ])
    zeilen.append(["Gesamt", fmt_euro(summen["gefordert"]),
                   fmt_euro(summen["gezahlt"]),
                   fmt_euro(summen["differenz"]) if summen["differenz"] > 0.005 else "–",
                   ""])
    erstelle_positions_tabelle(
        doc, ["Position", "gefordert", "gezahlt", "Differenz", "Anmerkung"],
        zeilen, spalten_breiten=[5.0, 2.6, 2.6, 2.6, 4.2])

    verlauf = [(z["datum"], p["label"], z["betrag"], z["versicherung"])
               for p in ueb["positionen"] for z in p["zahlungen"]]
    if verlauf:
        doc.add_paragraph()
        fuege_abschnittstitel_ein(doc, "Zahlungsverlauf")
        erstelle_positions_tabelle(
            doc, ["Datum", "Position", "Betrag", "Versicherung"],
            [[fmt_datum(d), lbl, fmt_euro(b), v]
             for d, lbl, b, v in sorted(verlauf)],
            spalten_breiten=[2.6, 6.0, 2.8, 5.6])

    doc.add_paragraph()
    fuege_abschnittstitel_ein(doc, "Ihre Anwaltskosten")
    ak = ueb["anwaltskosten"]
    if ak.get("gezahlt_von_gegner"):
        _absatz(doc, f"Unsere Gebühren in Höhe von "
                     f"{fmt_euro(ak['gezahlt_von_gegner'])} wurden von der "
                     f"Gegenseite getragen — für Sie kostenfrei.")
    elif ak.get("rvg_betrag"):
        _absatz(doc, f"Unsere Gebühren nach dem RVG in Höhe von "
                     f"{fmt_euro(ak['rvg_betrag'])} werden von der Gegenseite "
                     f"getragen — für Sie kostenfrei.")
    else:
        _absatz(doc, "Unsere Gebühren werden von der Gegenseite getragen — "
                     "für Sie kostenfrei.")

    schluss = ueb["schluss"]
    if schluss["text"]:
        doc.add_paragraph()
        fuege_abschnittstitel_ein(
            doc, "Abschluss" if modus == "abschluss" else "Ausblick")
        _absatz(doc, schluss["text"])
        if (schluss["typ"] == "vorbehalt_spaetfolgen"
                and schluss["verjaehrung_datum"]):
            _absatz(doc, f"Bitte beachten Sie: Ansprüche wegen etwaiger "
                         f"Spätfolgen verjähren am "
                         f"{fmt_datum(schluss['verjaehrung_datum'])}.",
                    bold=True)

    if ueb["bewertung_cta"]:
        _absatz(doc, "Wir würden uns freuen, wenn Sie Ihre Erfahrung mit "
                     "unserer Kanzlei in einer Google-Bewertung teilen.",
                size=9, farbe=GRAU)

    doc.add_paragraph()
    _absatz(doc, "Für Rückfragen stehen wir Ihnen gerne zur Verfügung.")
    doc.add_paragraph()
    _absatz(doc, "Mit freundlichen Grüßen")
    doc.add_paragraph()
    _absatz(doc, "Rechtsanwälte Koch, Schatz & Kollegen",
            bold=True, farbe=NAVY)

    fuege_fusszeile_ein(doc, az)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
