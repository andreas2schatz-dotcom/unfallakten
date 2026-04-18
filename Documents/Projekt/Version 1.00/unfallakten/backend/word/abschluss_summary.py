"""
Abschluss-Summary – "Das haben wir fuer Sie erreicht"
======================================================
Generiert ein DOCX bei Fallabschluss. Wird gespeichert als
dokumente.typ = 'sonstiges' mit portal_sichtbar = 1.
"""
import io
import logging
import sqlite3
from datetime import datetime, date

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

_POSITIONEN_LABELS = {
    "reparaturkosten":   "Reparaturkosten",
    "wiederbeschaffung": "Wiederbeschaffungswert",
    "wertminderung":     "Wertminderung",
    "nutzungsausfall":   "Nutzungsausfall",
    "mietwagenkosten":   "Mietwagenkosten",
    "sv_kosten":         "Sachverstaendigenkosten",
    "abschleppkosten":   "Abschleppkosten",
    "anabmeldekosten":   "An-/Abmeldekosten",
    "standkosten":       "Standkosten",
    "schmerzensgeld":    "Schmerzensgeld",
    "sonstiges":         "Sonstige Kosten",
}


def _fmt_euro(betrag):
    if betrag is None:
        return u"\u2013"
    return u"{:,.2f}\u00a0\u20ac".format(float(betrag)).replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt_datum(iso_str):
    if not iso_str:
        return u"\u2013"
    try:
        return datetime.strptime(iso_str[:10], "%Y-%m-%d").strftime("%d.%m.%Y")
    except ValueError:
        return iso_str


def generiere_abschluss_summary(conn, akte_id):
    # type: (sqlite3.Connection, str) -> bytes
    """Erstellt das DOCX als Bytes. Caller speichert auf Disk und in dokumente-Tabelle."""
    akte = conn.execute(
        "SELECT az, unfalldatum, status, haftungsquote, erstellt_am FROM unfallakte WHERE az = ?",
        (akte_id,)
    ).fetchone()
    if not akte:
        raise ValueError("Akte {!r} nicht gefunden".format(akte_id))

    mandant = conn.execute(
        "SELECT name, vorname FROM beteiligte WHERE akte_id = ? AND rolle = 'mandant' LIMIT 1",
        (akte_id,)
    ).fetchone()

    sp = conn.execute(
        "SELECT * FROM schadenpositionen WHERE akte_id = ?", (akte_id,)
    ).fetchone()

    reg_gesamt = conn.execute("""
        SELECT COALESCE(SUM(rp.betrag_reguliert), 0.0) AS total
        FROM regulierung_positionen rp
        JOIN abrechnungsschreiben ab ON ab.id = rp.abrechnungsschreiben_id
        WHERE ab.akte_id = ?
    """, (akte_id,)).fetchone()

    reg_per_pos = {
        r["position_key"]: float(r["reguliert"])
        for r in conn.execute("""
            SELECT rp.position_key, SUM(rp.betrag_reguliert) AS reguliert
            FROM regulierung_positionen rp
            JOIN abrechnungsschreiben ab ON ab.id = rp.abrechnungsschreiben_id
            WHERE ab.akte_id = ?
            GROUP BY rp.position_key
        """, (akte_id,)).fetchall()
    }

    heute = date.today()
    try:
        mandats_start = date.fromisoformat((akte["erstellt_am"] or "")[:10])
        dauer = (heute - mandats_start).days
    except (ValueError, TypeError):
        dauer = 0

    gesamt_gefordert = 0.0
    if sp:
        gesamt_gefordert = float(sum(
            sp[k] or 0.0
            for k in ("reparaturkosten", "wiederbeschaffung", "wertminderung",
                      "nutzungsausfall", "mietwagenkosten", "sv_kosten",
                      "abschleppkosten", "anabmeldekosten", "standkosten",
                      "schmerzensgeld", "sonstiges")
        ) - float(sp["restwert"] or 0.0))

    gesamt_reguliert = float(reg_gesamt["total"]) if reg_gesamt else 0.0
    quote = (gesamt_reguliert / gesamt_gefordert * 100) if gesamt_gefordert > 0 else 0.0

    doc = Document()

    h = doc.add_heading("Abschluss-Summary", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Kanzlei Koch, Schatz & Kollegen", 1)

    mandant_name = "Ihr Mandant"
    if mandant:
        mandant_name = u"{} {}".format(
            (mandant["vorname"] or "").strip(),
            (mandant["name"] or "").strip()
        ).strip()

    for label, wert in [
        ("Mandant", mandant_name),
        ("Aktenzeichen", akte["az"]),
        ("Unfalldatum", _fmt_datum(akte["unfalldatum"])),
        ("Verfahrensdauer", "{} Tage".format(dauer)),
    ]:
        p = doc.add_paragraph()
        p.add_run(u"{}: ".format(label)).bold = True
        p.add_run(wert)

    doc.add_heading("Das haben wir fuer Sie erreicht", 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(("Position", "Gefordert", "Erhalten")):
        r = hdr[i].paragraphs[0].add_run(txt)
        r.bold = True

    for key, label in _POSITIONEN_LABELS.items():
        if not sp:
            continue
        gef = float(sp[key] or 0.0)
        if key == "wiederbeschaffung":
            gef = max(0.0, gef - float(sp["restwert"] or 0.0))
        if gef <= 0:
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = _fmt_euro(gef)
        row_cells[2].text = _fmt_euro(reg_per_pos.get(key, 0.0))

    total_row = table.add_row().cells
    total_row[0].paragraphs[0].add_run("GESAMT").bold = True
    total_row[1].paragraphs[0].add_run(_fmt_euro(gesamt_gefordert)).bold = True
    total_row[2].paragraphs[0].add_run(_fmt_euro(gesamt_reguliert)).bold = True

    doc.add_paragraph()
    result_para = doc.add_paragraph()
    run = result_para.add_run(
        u"Wir konnten {} ({:.0f}\u00a0% Ihrer Forderung) fuer Sie durchsetzen.".format(
            _fmt_euro(gesamt_reguliert), quote
        )
    )
    run.bold = True
    run.font.size = Pt(13)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
