"""
Modul 8 – Tests: Wiedervorlage & Sachstandsanfragen
=====================================================
Tests ohne echte RA-Micro Verbindung (alles gemockt).

Testet:
  - Verbindungsstatus-Endpunkt (aktiv/deaktiviert)
  - Wiedervorlage-Liste (mit Mock-Daten)
  - Word-Generierung (Dokumentstruktur, Dateiname)
  - Sachstandsanfrage-Endpunkt (Download)
  - Fehlerbehandlung (RA-Micro nicht erreichbar)
"""

import io
import pytest
from unittest.mock import patch, MagicMock
from datetime import date, datetime


# ── Fixtures ──────────────────────────────────────────────────────────────────

MOCK_WV = {
    "GUIDWiedervorlage":       "55F4DD24-86AF-4A3D-95BE-FF65A3433EA9",
    "dtWiedervorlage":         datetime(2026, 3, 18),
    "sWiedervorlagegrund":     "Stellungnahme Gegner?schieben!",
    "wv_bemerkung":            "",
    "sachbearbeiter_kuerzel":  "TB",
    "akte_sachbearbeiter_kuerzel": "TB",
    "GUIDAkte":                "648D714B-791D-4BEC-AB9B-2BA66E69566E",
    "sAktenNummer":            "62260/25TB",
    "sAktenKurzBezeichnung":   "Müller ./. KRAVAG",
    "sAktenBezeichnung":       "Müller Hans ./. KRAVAG Versicherung",
    "sMandant":                "Müller, Hans",
    "sGegner":                 "KRAVAG Versicherung AG",
    "GUIDAdresse":             "AAAAAAAA-0000-0000-0000-000000000001",
    "sBeteiligtenKennzeichen": "HV",
    "sBetreffZeile1":          "Schadennummer KH-123456",
    "sBetreffZeile2":          "KH-Schaden vom 15.01.2026",
    "sBetreffZeile3":          "",
    "iAdressnummer":           99001,
    "sErsteAdresszeile":       "KRAVAG Versicherung AG",
    "adr_name":                "KRAVAG Versicherung AG",
    "adr_strasse":             "Schanzenstraße 85",
    "adr_plz":                 "20357",
    "adr_ort":                 "Hamburg",
    "adr_email":               "schaden@kravag.de",
    "adr_briefanrede":         "Sehr geehrte Damen und Herren,",
    "adr_telefon":             "040-1234567",
    "adr_telefax":             "040-1234568",
}


@pytest.fixture
def client(tmp_path, monkeypatch):
    """Test-Client mit deaktivierter RA-Micro Verbindung."""
    monkeypatch.setenv("RAMICRO_AKTIV", "false")
    monkeypatch.setenv("DB_PATH", str(tmp_path / "test.db"))
    monkeypatch.setenv("FLASK_SECRET_KEY", "test-secret")
    monkeypatch.setenv("JWT_SECRET_KEY", "test-jwt-secret")

    from backend.app import erstelle_app
    app = erstelle_app({"TESTING": True})
    return app.test_client()


@pytest.fixture
def auth_headers(client):
    """JWT-Token für Test-Anfragen."""
    r = client.post("/auth/login", json={
        "email": "koch@anwalt-offenbach.de",
        "passwort": "Kanzlei2024!"
    })
    if r.status_code != 200:
        pytest.skip("Seed-Daten nicht vorhanden – Auth-Login fehlgeschlagen")
    token = r.get_json().get("access_token", "")
    return {"Authorization": f"Bearer {token}"}


# ── Tests: Verbindungsstatus ──────────────────────────────────────────────────

class TestVerbindungStatus:

    def test_status_deaktiviert(self, client, auth_headers):
        """RAMICRO_AKTIV=false → Status 'deaktiviert'."""
        r = client.get("/wiedervorlage/status", headers=auth_headers)
        assert r.status_code == 200
        daten = r.get_json()
        assert daten["status"] == "deaktiviert"

    def test_status_ohne_auth(self, client):
        """Ohne Token → 401."""
        r = client.get("/wiedervorlage/status")
        assert r.status_code == 401

    def test_status_aktiv_verbindung_ok(self, client, auth_headers, monkeypatch):
        """RAMICRO_AKTIV=true + Mock → Status 'ok'."""
        monkeypatch.setenv("RAMICRO_AKTIV", "true")
        monkeypatch.setenv("RAMICRO_HOST", "192.168.1.1")
        monkeypatch.setenv("RAMICRO_USER", "test_user")

        with patch("backend.ramicro.connector.verbindung_pruefen",
                   return_value={"status": "ok", "host": "192.168.1.1", "datenbank": "RAMICRO"}):
            r = client.get("/wiedervorlage/status", headers=auth_headers)
        assert r.status_code == 200
        assert r.get_json()["status"] == "ok"

    def test_status_aktiv_verbindung_fehler(self, client, auth_headers, monkeypatch):
        """Verbindung fehlgeschlagen → 503."""
        with patch("backend.ramicro.connector.verbindung_pruefen",
                   return_value={"status": "fehler", "meldung": "Timeout"}):
            r = client.get("/wiedervorlage/status", headers=auth_headers)
        assert r.status_code == 503


# ── Tests: Wiedervorlage-Liste ────────────────────────────────────────────────

class TestWiedervorlageListe:

    def test_liste_ramicro_deaktiviert(self, client, auth_headers):
        """Deaktiviert → 503 mit Code RAMICRO_NICHT_AKTIV."""
        r = client.get("/wiedervorlage/", headers=auth_headers)
        assert r.status_code == 503
        assert r.get_json()["code"] == "RAMICRO_NICHT_AKTIV"

    def test_liste_mit_mock_daten(self, client, auth_headers):
        """Mock-Daten → korrekte Felder in Antwort."""
        with patch("backend.routers.wiedervorlage_routes.hole_faellige_wiedervorlagen",
                   return_value=[MOCK_WV]):
            r = client.get("/wiedervorlage/", headers=auth_headers)

        assert r.status_code == 200
        daten = r.get_json()
        assert daten["anzahl"] == 1
        wv = daten["wiedervorlagen"][0]

        assert wv["guid"]            == "55F4DD24-86AF-4A3D-95BE-FF65A3433EA9"
        assert wv["aktenzeichen"]    == "62260/25TB"
        assert wv["kurzbezeichnung"] == "Müller ./. KRAVAG"
        assert wv["gegner_hv_name"]  == "KRAVAG Versicherung AG"
        assert wv["gegner_hv_email"] == "schaden@kravag.de"
        assert wv["betreff1"]        == "Schadennummer KH-123456"
        assert wv["datum"]           == "2026-03-18"

    def test_liste_leer(self, client, auth_headers):
        """Keine WV fällig → anzahl=0, leere Liste."""
        with patch("backend.routers.wiedervorlage_routes.hole_faellige_wiedervorlagen",
                   return_value=[]):
            r = client.get("/wiedervorlage/", headers=auth_headers)
        assert r.status_code == 200
        assert r.get_json()["anzahl"] == 0

    def test_liste_filter_nur_heute(self, client, auth_headers):
        """?nur_heute=true wird an Service durchgereicht."""
        with patch("backend.routers.wiedervorlage_routes.hole_faellige_wiedervorlagen",
                   return_value=[]) as mock:
            client.get("/wiedervorlage/?nur_heute=true", headers=auth_headers)
            mock.assert_called_once_with(nur_heute=True, sachbearbeiter=None, limit=200)

    def test_liste_filter_sachbearbeiter(self, client, auth_headers):
        """?sb=AS wird an Service durchgereicht."""
        with patch("backend.routers.wiedervorlage_routes.hole_faellige_wiedervorlagen",
                   return_value=[]) as mock:
            client.get("/wiedervorlage/?sb=AS", headers=auth_headers)
            mock.assert_called_once_with(nur_heute=False, sachbearbeiter="AS", limit=200)

    def test_liste_verbindung_fehler(self, client, auth_headers):
        """VerbindungsFehler → 503."""
        from backend.ramicro.connector import RaMicroVerbindungsFehler
        with patch("backend.routers.wiedervorlage_routes.hole_faellige_wiedervorlagen",
                   side_effect=RaMicroVerbindungsFehler("Timeout")):
            r = client.get("/wiedervorlage/", headers=auth_headers)
        assert r.status_code == 503
        assert r.get_json()["code"] == "RAMICRO_VERBINDUNG_FEHLER"


# ── Tests: Word-Generierung ───────────────────────────────────────────────────

class TestWordGenerierung:

    def test_generiert_gueltige_docx(self):
        """generiere_sachstandsanfrage_wv() gibt gültige .docx-Bytes zurück."""
        from docx import Document
        from backend.word.sachstandsanfrage_wv import generiere_sachstandsanfrage_wv

        result = generiere_sachstandsanfrage_wv(MOCK_WV)

        assert isinstance(result, bytes)
        assert len(result) > 1000

        # Prüfen ob gültiges DOCX
        doc = Document(io.BytesIO(result))
        texte = [p.text for p in doc.paragraphs]
        text_gesamt = " ".join(texte)

        assert "KRAVAG Versicherung AG"              in text_gesamt
        assert "Schanzenstraße 85"                   in text_gesamt
        assert "20357 Hamburg"                       in text_gesamt
        assert "62260/25TB"                          in text_gesamt
        assert "Müller ./. KRAVAG"                   in text_gesamt
        assert "Schadennummer KH-123456"             in text_gesamt
        assert "Sehr geehrte Damen und Herren,"      in text_gesamt
        assert "Mit freundlichen Grüßen"             in text_gesamt

    def test_word_enthaelt_brieftext(self):
        """Standardbrief-Text ist im Dokument."""
        from docx import Document
        from backend.word.sachstandsanfrage_wv import generiere_sachstandsanfrage_wv

        result = generiere_sachstandsanfrage_wv(MOCK_WV)
        doc = Document(io.BytesIO(result))
        text_gesamt = " ".join(p.text for p in doc.paragraphs)

        assert "letztes Schreiben"       in text_gesamt
        assert "Sachstandsmitteilung"    in text_gesamt

    def test_dateiname_generierung(self):
        """dateiname_generieren() erzeugt korrekten Dateinamen."""
        from backend.word.sachstandsanfrage_wv import dateiname_generieren

        name = dateiname_generieren("62260/25TB", date(2026, 3, 18))
        assert name == "62260-25TB_sachstandsanfrage_2026-03-18.docx"
        # Keine Schrägstriche im Dateinamen
        assert "/" not in name

    def test_dateiname_ohne_datum(self):
        """Ohne Datum → heutiges Datum."""
        from backend.word.sachstandsanfrage_wv import dateiname_generieren
        name = dateiname_generieren("1213/25AS")
        assert date.today().isoformat() in name

    def test_datum_deutsch_formatierung(self):
        """Deutsches Datumsformat korrekt."""
        from backend.word.sachstandsanfrage_wv import _datum_deutsch
        assert _datum_deutsch(date(2026, 3, 12)) == "12. März 2026"
        assert _datum_deutsch(date(2026, 1, 1))  == "1. Januar 2026"
        assert _datum_deutsch(date(2026, 12, 31)) == "31. Dezember 2026"

    def test_word_mit_leerem_mandant(self):
        """Fehlende Adressdaten → kein Crash, Fallbacks greifen."""
        from backend.word.sachstandsanfrage_wv import generiere_sachstandsanfrage_wv
        wv_leer = {**MOCK_WV, "sErsteAdresszeile": "", "adr_name": "",
                   "adr_strasse": "", "adr_plz": "", "adr_ort": "",
                   "adr_email": "", "sBetreffZeile1": "", "sBetreffZeile2": ""}
        result = generiere_sachstandsanfrage_wv(wv_leer)
        assert isinstance(result, bytes)
        assert len(result) > 500


# ── Tests: Sachstandsanfrage-Endpunkt ────────────────────────────────────────

class TestSachstandsanfrageEndpunkt:

    def test_download_erfolgreich(self, client, auth_headers):
        """POST mit Mock-Daten → DOCX-Download."""
        with patch("backend.routers.wiedervorlage_routes.hole_wiedervorlage_details",
                   return_value=MOCK_WV), \
             patch("backend.routers.wiedervorlage_routes.logge_aktivitaet"):

            r = client.post(
                "/wiedervorlage/55F4DD24-86AF-4A3D-95BE-FF65A3433EA9/sachstandsanfrage",
                headers=auth_headers
            )

        assert r.status_code == 200
        assert "wordprocessingml" in r.content_type
        assert len(r.data) > 1000

    def test_download_nicht_gefunden(self, client, auth_headers):
        """WV nicht in RA-Micro → 404."""
        with patch("backend.routers.wiedervorlage_routes.hole_wiedervorlage_details",
                   return_value=None):
            r = client.post(
                "/wiedervorlage/NICHT-VORHANDEN/sachstandsanfrage",
                headers=auth_headers
            )
        assert r.status_code == 404

    def test_download_ramicro_fehler(self, client, auth_headers):
        """RA-Micro nicht erreichbar → 503."""
        from backend.ramicro.connector import RaMicroVerbindungsFehler
        with patch("backend.routers.wiedervorlage_routes.hole_wiedervorlage_details",
                   side_effect=RaMicroVerbindungsFehler("Timeout")):
            r = client.post(
                "/wiedervorlage/55F4DD24/sachstandsanfrage",
                headers=auth_headers
            )
        assert r.status_code == 503

    def test_aktivitaet_wird_geloggt(self, client, auth_headers):
        """Nach erfolgreichem Download → Aktivitätslog-Aufruf."""
        with patch("backend.routers.wiedervorlage_routes.hole_wiedervorlage_details",
                   return_value=MOCK_WV), \
             patch("backend.routers.wiedervorlage_routes.logge_aktivitaet") as mock_log:

            client.post(
                "/wiedervorlage/55F4DD24-86AF-4A3D-95BE-FF65A3433EA9/sachstandsanfrage",
                headers=auth_headers
            )

        mock_log.assert_called_once()
        args = mock_log.call_args
        assert args.kwargs["aktion"] == "sachstandsanfrage_wv"
        assert "62260/25TB" in args.kwargs["beschreibung"]


# ── Tests: Sachbearbeiter-Mapping ─────────────────────────────────────────────

class TestSachbearbeiter:

    def test_bekanntes_kuerzel(self):
        from backend.ramicro.sachbearbeiter import hole_sachbearbeiter
        sb = hole_sachbearbeiter("AS")
        assert sb["name"] == "Andreas Schatz"
        assert sb["titel"] == "Rechtsanwalt"

    def test_unbekanntes_kuerzel(self):
        from backend.ramicro.sachbearbeiter import hole_sachbearbeiter
        sb = hole_sachbearbeiter("XX")
        assert "[XX]" in sb["name"]

    def test_leeres_kuerzel(self):
        from backend.ramicro.sachbearbeiter import hole_sachbearbeiter
        sb = hole_sachbearbeiter("")
        assert "Koch" in sb["name"] or sb["name"]  # Fallback-Text

    def test_kleinschreibung(self):
        from backend.ramicro.sachbearbeiter import hole_sachbearbeiter
        sb = hole_sachbearbeiter("as")
        assert sb["name"] == "Andreas Schatz"


# ── Tests: Statistik ──────────────────────────────────────────────────────────

class TestStatistik:

    def test_statistik_deaktiviert(self, client, auth_headers):
        """RA-Micro deaktiviert → 503."""
        r = client.get("/wiedervorlage/statistik", headers=auth_headers)
        assert r.status_code == 503

    def test_statistik_mit_mock(self, client, auth_headers):
        """Mock-Daten → Gruppen in Antwort."""
        mock_data = {"gruppen": [
            {"sWiedervorlagegrund": "Stellungnahme Gegner", "anzahl": 5}
        ]}
        with patch("backend.routers.wiedervorlage_routes.hole_wiedervorlagen_statistik",
                   return_value=mock_data):
            r = client.get("/wiedervorlage/statistik", headers=auth_headers)
        assert r.status_code == 200
        assert len(r.get_json()["gruppen"]) == 1
