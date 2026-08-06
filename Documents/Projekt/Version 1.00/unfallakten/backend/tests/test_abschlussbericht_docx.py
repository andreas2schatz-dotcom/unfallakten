"""DOCX-Smoke-Tests für den Abschluss-/Sachstandsbericht."""
import io
import os
import sys
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

from docx import Document

from backend.word.abschlussbericht import generiere_abschlussbericht


def _daten(schluss_typ="endgueltig"):
    return {
        "akte": {"aktenzeichen": "42/26", "unfalldatum": "2026-01-10",
                 "unfallort": "Offenbach", "haftungsquote": 100.0},
        "mandant": {"name": "Muster", "vorname": "Max", "anrede": "1",
                    "anschrift": "Weg 1", "plz": "63065", "ort": "Offenbach",
                    "vorsteuer": "N"},
        "gegner": {"versicherung": "HUK-COBURG"},
        "schaden": {"nutzungsausfall": 300.0, "mietwagenkosten": 500.0},
        "abrechnungen": [{
            "datum": "2026-02-01", "versicherung": "HUK-COBURG",
            "gesamt_reguliert": 650.0, "haftungsquote": 100.0,
            "positionen": [
                {"position_key": "nutzungsausfall",
                 "betrag_gefordert": 300.0, "betrag_reguliert": 300.0},
                {"position_key": "mietwagenkosten",
                 "betrag_gefordert": 500.0, "betrag_reguliert": 350.0,
                 "kuerzungsart_bezeichnung": "Überhöhter Tagessatz"}],
        }],
        "wdm_roh": {},
        "abschluss_status": {"schluss_typ": schluss_typ,
                             "schluss_text": "Damit ist die Sache erledigt.",
                             "naechste_schritte_text": "Wir warten auf die HUK."},
        "gebuehren_kontext": {"faktor": 1.3, "streitwert": 800.0,
                              "erstellt_am": "2026-01-15"},
        "kanzlei": None,
    }


def _volltext(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    teile = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                teile.append(cell.text)
    return "\n".join(teile)


class TestAbschlussberichtDocx(unittest.TestCase):

    def test_abschluss_variante(self):
        b = generiere_abschlussbericht(_daten("endgueltig"))
        self.assertGreater(len(b), 5000)
        text = _volltext(b)
        self.assertIn("42/26", text)
        self.assertIn("Abschluss", text)
        self.assertIn("650,00", text)
        self.assertIn("Überhöhter Tagessatz", text)
        self.assertIn("Damit ist die Sache erledigt.", text)
        self.assertIn("Mit freundlichen Grüßen", text)
        self.assertIn("Koch, Schatz", text)

    def test_sachstand_variante(self):
        b = generiere_abschlussbericht(_daten("offen"))
        text = _volltext(b)
        self.assertIn("Sachstandsbericht", text)
        self.assertIn("Wir warten auf die HUK.", text)
        self.assertNotIn("Für Sie durchgesetzt", text)

    def test_teilhaftung_mit_kostenfrei_aussage(self):
        daten = _daten("endgueltig")
        daten["abrechnungen"][0]["haftungsquote"] = 70.0
        b = generiere_abschlussbericht(daten)
        text = _volltext(b)
        self.assertIn("kostenfrei", text)
        self.assertIn("regulierten Betrag", text)
        self.assertNotIn("informieren wir Sie gesondert", text)

    def test_sachstand_ohne_kuratiertes_feld(self):
        daten = _daten()
        daten["abschluss_status"] = None
        b = generiere_abschlussbericht(daten)
        self.assertIn("Sachstandsbericht", _volltext(b))

    def test_verjaehrungshinweis_auch_ohne_schlusstext(self):
        daten = _daten("vorbehalt_spaetfolgen")
        daten["abschluss_status"]["schluss_text"] = ""
        daten["abschluss_status"]["verjaehrung_datum"] = "2029-06-30"
        text = _volltext(generiere_abschlussbericht(daten))
        self.assertIn("verjähren am", text)
        self.assertIn("30.06.2029", text)


if __name__ == "__main__":
    unittest.main()
